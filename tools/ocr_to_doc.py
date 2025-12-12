# tools/ocr_to_doc.py
import sys
import os
from pathlib import Path
from typing import List, Tuple
from docx import Document
from openpyxl import Workbook
import flet as ft
from PIL import Image

# 初始化 OCR（只初始化一次）
_ocr_engine = None

# 延迟导入PaddleOCR，避免启动时加载
PaddleOCR = None

def get_ocr_engine():
    global _ocr_engine, PaddleOCR
    
    # 延迟导入PaddleOCR，只有在第一次使用时才导入
    if PaddleOCR is None:
        from paddleocr import PaddleOCR as _PaddleOCR
        PaddleOCR = _PaddleOCR
    
    if _ocr_engine is None:
        # 获取当前脚本所在目录
        script_dir = Path(__file__).parent.parent
        # 为PaddleOCR设置模型缓存路径到应用目录，避免打包后权限问题
        model_dir = script_dir / "paddleocr_models"
        model_dir.mkdir(exist_ok=True)
        
        # use_angle_cls=True 启用方向分类，更准
        _ocr_engine = PaddleOCR(
            use_angle_cls=True,
            lang="ch",  # 中文+英文
            # 设置模型缓存目录，避免打包后权限问题
            det_model_dir=str(model_dir / "det"),
            rec_model_dir=str(model_dir / "rec"),
            cls_model_dir=str(model_dir / "cls")
            # show_log和use_gpu参数已在新版本中移除
        )
    return _ocr_engine


def ocr_image_or_pdf(file_path: Path) -> str:
    """对单张图片或 PDF（转图）进行 OCR，返回纯文本（按行拼接）"""
    ocr = get_ocr_engine()
    results = []

    if file_path.suffix.lower() == ".pdf":
        from pdf2image import convert_from_path
        # 复用你已有的 poppler 路径逻辑（如果需要）
        images = convert_from_path(str(file_path), dpi=150)
    else:
        images = [Image.open(file_path)]

    for img in images:
        # PaddleOCR 接受 PIL.Image 或 numpy array
        ocr_result = ocr.ocr(img, cls=True)
        if not ocr_result or not ocr_result[0]:
            continue
        # 提取文本（忽略坐标和置信度）
        text_lines = [line[1][0] for line in ocr_result[0]]
        results.append("\n".join(text_lines))

    return "\n---分页---\n".join(results) if len(results) > 1 else (results[0] if results else "")


def export_to_word(texts: List[Tuple[str, str]], output_path: Path):
    """texts: [(filename, content), ...]"""
    doc = Document()
    doc.add_heading("OCR 识别结果", 0)

    for filename, content in texts:
        doc.add_heading(f"📄 {filename}", level=1)
        doc.add_paragraph(content)
        doc.add_page_break()

    doc.save(output_path)


def export_to_excel(texts: List[Tuple[str, str]], output_path: Path):
    """每行一个文件，A列文件名，B列内容"""
    wb = Workbook()
    ws = wb.active
    ws.title = "OCR Results"
    ws.append(["文件名", "识别内容"])

    for filename, content in texts:
        # Excel 单元格有字符限制（32767），长文本可截断或警告
        ws.append([filename, content[:32000]])

    wb.save(output_path)


def collect_images_or_pdfs(input_path: Path) -> List[Path]:
    allowed_ext = {".jpg", ".jpeg", ".png", ".bmp", ".pdf"}
    if input_path.is_file() and input_path.suffix.lower() in allowed_ext:
        return [input_path]
    elif input_path.is_dir():
        return sorted([
            p for p in input_path.rglob("*")
            if p.suffix.lower() in allowed_ext
        ])
    return []


def create_ocr_tool_page(page: ft.Page) -> ft.Control:
    input_field = ft.TextField(label="输入路径（图片或PDF）", read_only=True, width=400)
    output_field = ft.TextField(label="输出目录", read_only=True, width=400)
    format_dropdown = ft.Dropdown(
        label="输出格式",
        options=[
            ft.dropdown.Option("word", "Word (.docx)"),
            ft.dropdown.Option("excel", "Excel (.xlsx)"),
        ],
        value="word",
        width=200
    )
    status_text = ft.Text("", size=13, selectable=True, expand=True)

    # File pickers
    file_picker = ft.FilePicker()
    folder_picker_in = ft.FilePicker()
    folder_picker_out = ft.FilePicker()
    page.overlay.extend([file_picker, folder_picker_in, folder_picker_out])

    def on_input(e):
        path = e.path or (e.files[0].path if e.files else None)
        if path:
            input_field.value = path
            input_field.update()

    def on_output(e):
        if e.path:
            output_field.value = e.path
            output_field.update()

    file_picker.on_result = on_input
    folder_picker_in.on_result = on_input
    folder_picker_out.on_result = on_output

    def start_ocr(_):
        input_str = input_field.value
        output_str = output_field.value
        fmt = format_dropdown.value

        if not input_str or not os.path.exists(input_str):
            status_text.value = "❌ 请输入有效的输入路径"
            status_text.color = "red"
            status_text.update()
            return
        if not output_str:
            status_text.value = "❌ 请选择输出目录"
            status_text.color = "red"
            status_text.update()
            return

        input_p = Path(input_str)
        output_p = Path(output_str)
        files = collect_images_or_pdfs(input_p)

        if not files:
            status_text.value = "❌ 未找到支持的图片或PDF文件"
            status_text.color = "red"
            status_text.update()
            return

        status_text.value = f"🔄 正在识别 {len(files)} 个文件...\n"
        status_text.color = "blue"
        status_text.update()

        results = []
        for f in files:
            try:
                text = ocr_image_or_pdf(f)
                results.append((f.name, text))
                status_text.value += f"✅ {f.name} 识别完成\n"
            except Exception as e:
                status_text.value += f"❌ {f.name} 失败: {str(e)[:100]}\n"
            status_text.update()

        # 导出
        if fmt == "word":
            out_file = output_p / "OCR结果.docx"
            export_to_word(results, out_file)
        else:
            out_file = output_p / "OCR结果.xlsx"
            export_to_excel(results, out_file)

        status_text.value += f"\n🎉 完成！文件已保存至:\n{out_file}"
        status_text.color = "green"
        status_text.update()

        # 自动打开文件夹
        try:
            os.startfile(output_p)
        except:
            pass

    return ft.Column([
        ft.Text("🔍 OCR 文字识别（图片/PDF → Word/Excel）", size=24, weight="bold"),
        ft.Row([
            ft.Column([
                ft.Text("📥 输入:", weight="bold"),
                input_field,
                ft.Row([
                    ft.ElevatedButton("选择文件", icon=ft.Icons.FILE_PRESENT, on_click=lambda _: file_picker.pick_files(allowed_extensions=["jpg","jpeg","png","bmp","pdf"])),
                    ft.ElevatedButton("选择文件夹", icon=ft.Icons.FOLDER_OPEN, on_click=lambda _: folder_picker_in.get_directory_path()),
                ])
            ]),
            ft.VerticalDivider(),
            ft.Column([
                ft.Text("📤 输出:", weight="bold"),
                output_field,
                ft.ElevatedButton("选择输出目录", icon=ft.Icons.SAVE, on_click=lambda _: folder_picker_out.get_directory_path()),
                format_dropdown,
            ])
        ]),
        ft.Divider(height=25),
        ft.ElevatedButton("开始识别", icon=ft.Icons.PLAY_ARROW, on_click=start_ocr, height=50, style=ft.ButtonStyle(bgcolor=ft.colors.GREEN, color=ft.colors.WHITE)),
        ft.Divider(),
        ft.Container(content=status_text, padding=10, border=ft.border.all(1, ft.colors.GREY_300), border_radius=8, bgcolor=ft.colors.BLACK12, expand=True)
    ], expand=True, scroll=ft.ScrollMode.AUTO)


# 注册工具
from . import register_tool
register_tool("OCR文字识别", ft.Icons.TEXT_SNIPPET, create_ocr_tool_page)
